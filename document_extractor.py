from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
import threading
from typing import Dict, List, Tuple, Set, Optional
import fitz  # PyMuPDF
from docx import Document
import io
from PIL import Image
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from tqdm import tqdm
import cProfile
import time
import os
import xml.etree.ElementTree as etree
from lxml import etree
import logging
import base64
import uuid
from pathlib import Path

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentExtractor:
    def __init__(self):
        self.supported_formats = ['.docx']
        # 公式相关命名空间
        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'v': 'urn:schemas-microsoft-com:vml'
        }
        
        # 支持的图片格式
        self.image_formats = {
            'image/png': '.png',
            'image/jpeg': '.jpg',
            'image/jpg': '.jpg',
            'image/gif': '.gif',
            'image/bmp': '.bmp',
            'image/tiff': '.tiff',
            'image/x-emf': '.emf',
            'image/x-wmf': '.wmf',
            'image/x-pict': '.pict'
        }
        
        # 数学符号和LaTeX标记的映射
        self.math_symbols = {
            # 希腊字母
            'α': '\\alpha', 'β': '\\beta', 'γ': '\\gamma', 'δ': '\\delta',
            'ε': '\\epsilon', 'ζ': '\\zeta', 'η': '\\eta', 'θ': '\\theta',
            'ι': '\\iota', 'κ': '\\kappa', 'λ': '\\lambda', 'μ': '\\mu',
            'ν': '\\nu', 'ξ': '\\xi', 'ο': 'o', 'π': '\\pi',
            'ρ': '\\rho', 'σ': '\\sigma', 'τ': '\\tau', 'υ': '\\upsilon',
            'φ': '\\phi', 'χ': '\\chi', 'ψ': '\\psi', 'ω': '\\omega',
            
            # 数学运算符
            '±': '\\pm', '×': '\\times', '÷': '\\div', '∑': '\\sum',
            '∏': '\\prod', '∫': '\\int', '∂': '\\partial', '∞': '\\infty',
            '≠': '\\neq', '≤': '\\leq', '≥': '\\geq', '≈': '\\approx',
        }

    def extract_content(self, file_path: str) -> Dict:
        """提取文档内容"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.docx':
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _extract_from_docx(self, file_path: str) -> Dict[str, List]:
        """从DOCX文件提取内容"""
        logger.info("开始处理DOCX文件...")
        start_time = time.time()
        
        doc = Document(file_path)
        content_blocks = []
        current_index = 0
        
        # 创建图片输出目录
        output_dir = os.path.abspath(os.path.join("output", "images"))
        os.makedirs(output_dir, exist_ok=True)
        
        # 预先收集所有图像关系
        image_rels = {}
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                image_rels[rel_id] = rel
        
        logger.info(f"文档中包含 {len(image_rels)} 个图像关系")

        # 获取文档主体部分
        body = doc.element.body
        
        # 遍历所有子元素（包括段落和其他元素）
        for element in body.iterchildren():
            element_tag = element.tag.split('}')[-1]  # 获取标签名
            logger.info(f"处理元素: {element_tag}")

            if element_tag == 'p':  # 段落
                # 创建python-docx的段落对象
                paragraph = doc.paragraphs[current_index]
                
                # 检查是否为空段落
                if not paragraph.text.strip():  # 使用 text 替代 text_content()
                    # 检查是否包含图片
                    images = self._process_images(paragraph, output_dir, image_rels)
                    for image in images:
                        image['original_index'] = current_index
                        content_blocks.append(image)
                    
                    # 检查是否包含公式
                    formula = self._extract_formula(paragraph._element)
                    if formula:
                        formula['original_index'] = current_index
                        content_blocks.append(formula)
                    
                    current_index += 1
                    continue

                # 获取段落的基本格式信息
                format_info = {
                    'style_name': paragraph.style.name,
                    'alignment': str(paragraph.alignment),
                    'line_spacing': paragraph.paragraph_format.line_spacing,
                    'first_line_indent': paragraph.paragraph_format.first_line_indent
                }

                # 检查段落是否包含公式
                formula = self._extract_formula(paragraph._element)
                if formula:
                    formula['original_index'] = current_index
                    formula['format_info'] = format_info
                    content_blocks.append(formula)
                    current_index += 1
                    continue
                
                # 处理标题
                if paragraph.style.name.startswith('Heading'):
                    try:
                        level = int(paragraph.style.name.replace('Heading ', ''))
                        content_blocks.append({
                            'type': 'heading',
                            'content': paragraph.text.strip(),
                            'level': level,
                            'style': {
                                'is_heading': True,
                                'heading_level': level
                            },
                            'format_info': format_info,
                            'original_index': current_index
                        })
                    except ValueError:
                        content_blocks.append({
                            'type': 'text',
                            'content': paragraph.text.strip(),
                            'format_info': format_info,
                            'original_index': current_index
                        })
                else:
                    # 处理普通文本段落
                    content_blocks.append({
                        'type': 'text',
                        'content': paragraph.text.strip(),
                        'format_info': format_info,
                        'original_index': current_index
                    })

                # 处理段落中的图片
                images = self._process_images(paragraph, output_dir, image_rels)
                for image in images:
                    image['original_index'] = current_index
                    content_blocks.append(image)
                
                current_index += 1

            elif element_tag in ['drawing', 'pict']:  # 独立的图片元素
                logger.info(f"发现独立图片元素")
                # 创建一个临时段落对象来处理图片
                temp_para = doc.add_paragraph()
                temp_para._element = element
                images = self._process_images(temp_para, output_dir, image_rels)
                for image in images:
                    image['original_index'] = current_index
                    content_blocks.append(image)
                    current_index += 1
                # 删除临时段落
                temp_para._element.getparent().remove(temp_para._element)
                
            elif element_tag == 'oMath' or element_tag == 'oMathPara':  # 独立公式元素
                logger.info(f"发现独立公式元素")
                formula = self._extract_formula(element)
                if formula:
                    formula['original_index'] = current_index
                    content_blocks.append(formula)
                    current_index += 1

        logger.info(f"DOCX处理完成，耗时: {time.time() - start_time:.2f}秒")
        
        # 生成文档结构（仅包含标题）
        structure = [
            {'level': block['level'], 'text': block['content']}
            for block in content_blocks
            if block['type'] == 'heading'
        ]
        
        return {
            'structure': structure,
            'content_blocks': content_blocks
        }

    def _extract_formula(self, element) -> Optional[Dict]:
        """提取段落中的数学公式
        
        Returns:
            Dict: 包含公式信息的字典，如果没有公式则返回None
        """
        # 检查是否包含Office Math Markup Language (OMML) 公式
        o_math = element.find('.//m:oMath', namespaces=self.namespaces)
        if o_math is not None:
            logger.info("发现Office数学公式 (OMML)")
            # 提取公式文本
            formula_text = self._extract_omml_formula(o_math)
            if formula_text:
                return {
                    'type': 'formula',
                    'content': formula_text,
                    'is_formula': True,
                    'formula_type': 'omml'
                }
        
        # 检查是否包含LaTeX风格的公式
        text_content = ''.join([t.text for t in element.findall('.//w:t', namespaces=self.namespaces) if t.text])
        if self._contains_latex_formula(text_content):
            logger.info("发现LaTeX风格公式")
            return {
                'type': 'formula',
                'content': text_content,
                'is_formula': True,
                'formula_type': 'latex'
            }
        
        return None
    
    def _contains_latex_formula(self, text: str) -> bool:
        """检查文本是否包含LaTeX公式"""
        # 常见的LaTeX公式开始和结束标记
        latex_markers = [
            ('\\begin{equation}', '\\end{equation}'),
            ('\\begin{align}', '\\end{align}'),
            ('\\[', '\\]'),
            ('$$', '$$'),
            ('$', '$')
        ]
        
        # 检查是否包含成对的LaTeX标记
        for start, end in latex_markers:
            if start in text and end in text:
                start_pos = text.find(start)
                end_pos = text.find(end, start_pos + len(start))
                if end_pos > start_pos:
                    return True
        
        # 检查是否包含常见的LaTeX命令和符号组合
        latex_commands = ['\\frac', '\\sum', '\\int', '\\prod', '\\alpha', '\\beta', 
                          '\\Delta', '\\partial', '\\infty', '\\in', '\\subset']
        
        command_count = sum(1 for cmd in latex_commands if cmd in text)
        # 如果包含多个LaTeX命令，很可能是公式
        if command_count >= 2:
            return True
            
        return False
    
    def _extract_omml_formula(self, o_math_element) -> str:
        """从Office Math Markup Language (OMML)元素中提取公式文本，并尝试转换为LaTeX格式
        
        Args:
            o_math_element: OMML元素
            
        Returns:
            str: 提取的公式文本，尽可能转换为LaTeX格式
        """
        # 尝试从OMML转换为LaTeX
        try:
            # 保留原始OMML XML以确保兼容性
            xml_str = etree.tostring(o_math_element, encoding='unicode')
            
            # 简单的OMML到LaTeX的转换
            # 这里只做基本转换，完整转换需要更复杂的解析
            
            # 分数
            fractions = o_math_element.findall('.//m:f', namespaces=self.namespaces)
            if fractions:
                logger.info(f"公式中包含{len(fractions)}个分数")
                
            # 获取所有文本内容
            texts = o_math_element.findall('.//m:t', namespaces=self.namespaces)
            formula_text = ''.join([t.text if t.text else '' for t in texts])
            
            # 替换数学符号为LaTeX命令
            for symbol, latex_cmd in self.math_symbols.items():
                formula_text = formula_text.replace(symbol, latex_cmd)
            
            logger.info(f"提取的公式: {formula_text}")
            return formula_text
        except Exception as e:
            logger.error(f"公式提取失败: {str(e)}")
            # 失败时返回原始XML字符串
            return etree.tostring(o_math_element, encoding='unicode')

    def _process_images(self, paragraph, output_dir: str, image_rels: Dict) -> List[Dict]:
        """处理段落中的图片，改进版本，防止图片被截断或不完整
        
        Args:
            paragraph: 段落对象
            output_dir: 图片输出目录
            image_rels: 文档中的图片关系字典
            
        Returns:
            List[Dict]: 图片信息列表
        """
        results = []
        processed_rids = set()
        
        logger.info(f"处理段落图片: {paragraph.text[:30]}...")
        
        try:
            # 首先尝试使用新方法处理图片 (针对Word 2010及以上版本)
            drawings = self._find_all_drawings(paragraph._element)
            if drawings:
                for drawing_info in drawings:
                    if drawing_info['rid'] and drawing_info['rid'] not in processed_rids:
                        try:
                            # 获取图片数据
                            rid = drawing_info['rid']
                            if rid in paragraph.part.rels:
                                image_part = paragraph.part.rels[rid].target_part
                                
                                # 获取内容类型和扩展名
                                content_type = image_part.content_type
                                ext = self.image_formats.get(content_type, '.png')
                                
                                # 生成唯一文件名
                                unique_id = str(uuid.uuid4())[:8]
                                image_filename = f"image_{unique_id}{ext}"
                                image_path = os.path.join(output_dir, image_filename)
                                
                                # 保存图片数据 - 以二进制模式直接写入，保持原始格式和质量
                                with open(image_path, 'wb') as f:
                                    f.write(image_part.blob)
                                
                                # 尝试使用PIL验证图片完整性
                                self._validate_and_fix_image(image_path)
                                
                                logger.info(f"保存图片: {image_filename} (大小: {len(image_part.blob)} 字节)")
                                
                                # 构建图片信息对象
                                image_info = {
                                    'type': 'image',
                                    'image_path': image_path,
                                    'image_filename': image_filename,
                                    'position_info': drawing_info['position'],
                                    'file_size': len(image_part.blob),
                                    'content_type': content_type
                                }
                                results.append(image_info)
                                processed_rids.add(rid)
                        except Exception as e:
                            logger.error(f"处理图片出错 (新方法): {str(e)}")
            
            # 如果没有找到图片或处理失败，尝试旧方法 (兼容Word 2007)
            if not results:
                # 检查VML图片 (兼容旧版Word)
                shapes = paragraph._element.findall('.//v:shape', {'v': 'urn:schemas-microsoft-com:vml'})
                for shape in shapes:
                    image_data = shape.find('.//v:imagedata', {'v': 'urn:schemas-microsoft-com:vml'})
                    if image_data is not None:
                        rid = image_data.get(qn('r:id')) or image_data.get(qn('o:relid'))
                        if rid and rid not in processed_rids:
                            try:
                                # 获取图片数据
                                image_part = paragraph.part.rels[rid].target_part
                                
                                # 基于内容类型确定扩展名
                                content_type = image_part.content_type
                                ext = self.image_formats.get(content_type, '.png')
                                
                                # 生成唯一文件名
                                unique_id = str(uuid.uuid4())[:8]
                                image_filename = f"image_{unique_id}{ext}"
                                image_path = os.path.join(output_dir, image_filename)
                                
                                # 保存图片数据
                                with open(image_path, 'wb') as f:
                                    f.write(image_part.blob)
                                
                                # 验证图片
                                self._validate_and_fix_image(image_path)
                                
                                # 获取shape样式信息
                                position_info = {
                                    'type': 'shape',
                                    'style': shape.get('style', '')
                                }
                                
                                # 添加图片信息到结果列表
                                image_info = {
                                    'type': 'image',
                                    'image_path': image_path,
                                    'image_filename': image_filename,
                                    'position_info': position_info,
                                    'file_size': len(image_part.blob),
                                    'content_type': content_type
                                }
                                results.append(image_info)
                                processed_rids.add(rid)
                            except Exception as e:
                                logger.error(f"处理图片出错 (VML方法): {str(e)}")
            
            # 如果仍未找到图片，尝试直接查找所有的图片关系
            if not results:
                for rel_id, rel in paragraph.part.rels.items():
                    if "image" in rel.reltype and rel_id not in processed_rids:
                        try:
                            # 获取图片数据
                            image_part = rel.target_part
                            
                            # 基于内容类型确定扩展名
                            content_type = image_part.content_type
                            ext = self.image_formats.get(content_type, '.png')
                            
                            # 生成唯一文件名
                            unique_id = str(uuid.uuid4())[:8]
                            image_filename = f"image_{unique_id}{ext}"
                            image_path = os.path.join(output_dir, image_filename)
                            
                            # 保存图片数据
                            with open(image_path, 'wb') as f:
                                f.write(image_part.blob)
                            
                            # 验证图片
                            self._validate_and_fix_image(image_path)
                            
                            # 添加图片信息到结果列表
                            image_info = {
                                'type': 'image',
                                'image_path': image_path,
                                'image_filename': image_filename,
                                'position_info': {'type': 'unknown'},
                                'file_size': len(image_part.blob),
                                'content_type': content_type
                            }
                            results.append(image_info)
                            processed_rids.add(rel_id)
                        except Exception as e:
                            logger.error(f"处理图片关系出错: {str(e)}")
        
        except Exception as e:
            logger.error(f"处理段落图片时出错: {str(e)}")
        
        return results
    
    def _validate_and_fix_image(self, image_path: str) -> bool:
        """验证图片完整性，如有必要进行修复
        
        Args:
            image_path: 图片文件路径
            
        Returns:
            bool: 图片是否有效
        """
        try:
            # 尝试打开和验证图片
            with Image.open(image_path) as img:
                # 验证图片可以加载
                img.verify()
                return True
        except Exception as e:
            logger.warning(f"图片验证失败: {str(e)}")
            
            try:
                # 尝试进行修复 - 重新保存图片
                with Image.open(image_path) as img:
                    # 转换为RGB模式，解决一些颜色模式问题
                    if img.mode not in ('RGB', 'RGBA'):
                        img = img.convert('RGB')
                    
                    # 保存为PNG格式，通常更可靠
                    new_path = f"{os.path.splitext(image_path)[0]}_fixed.png"
                    img.save(new_path, format='PNG')
                    
                    # 如果修复成功，替换原文件
                    os.replace(new_path, image_path)
                    logger.info(f"图片已修复: {image_path}")
                    return True
            except Exception as repair_e:
                logger.error(f"图片修复失败: {str(repair_e)}")
                return False

    def _find_all_drawings(self, element) -> List[Dict]:
        """查找元素中的所有绘图元素，包括内联和浮动图片
        
        Args:
            element: XML元素
            
        Returns:
            List[Dict]: 绘图信息列表，每个包含rid和位置信息
        """
        results = []
        
        # 查找内联和浮动绘图
        for drawing_type in ['inline', 'anchor']:
            xpath = f'.//wp:drawing/wp:{drawing_type}'
            drawings = element.findall(xpath, namespaces=self.namespaces)
            
            for drawing in drawings:
                # 查找blip元素 (实际图片引用)
                blip = drawing.find('.//a:blip', namespaces=self.namespaces)
                if blip is not None:
                    # 获取图片关系ID
                    r_embed = blip.get(qn('r:embed'))
                    r_link = blip.get(qn('r:link'))
                    rid = r_embed or r_link
                    
                    if rid:
                        # 获取位置信息
                        position_info = {'type': drawing_type}
                        
                        if drawing_type == 'inline':
                            # 获取内联图片尺寸
                            extent = drawing.find('.//wp:extent', namespaces=self.namespaces)
                            if extent is not None:
                                position_info['width'] = extent.get('cx')
                                position_info['height'] = extent.get('cy')
                        else:
                            # 获取浮动图片位置
                            position_h = drawing.find('.//wp:positionH', namespaces=self.namespaces)
                            position_v = drawing.find('.//wp:positionV', namespaces=self.namespaces)
                            
                            if position_h is not None:
                                position_info['position_h'] = position_h.get('relativeFrom')
                                position_info['align_h'] = position_h.findtext('.//wp:align', namespaces=self.namespaces)
                                position_info['posOffset_h'] = position_h.findtext('.//wp:posOffset', namespaces=self.namespaces)
                            
                            if position_v is not None:
                                position_info['position_v'] = position_v.get('relativeFrom')
                                position_info['align_v'] = position_v.findtext('.//wp:align', namespaces=self.namespaces)
                                position_info['posOffset_v'] = position_v.findtext('.//wp:posOffset', namespaces=self.namespaces)
                        
                        results.append({
                            'rid': rid,
                            'position': position_info
                        })
        
        return results 