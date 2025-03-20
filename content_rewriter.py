from document_extractor import DocumentExtractor
from zhipu_api import ZhipuAPI
from docx import Document
from docx.shared import Inches
import os
import markdown
import shutil
import multiprocessing
import hashlib
from typing import Dict, List, Callable, Optional, Tuple, Any
from concurrent.futures import ThreadPoolExecutor, as_completed
import tqdm
import time
import json
import logging
import psutil  # 需要添加到requirements.txt

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ContentRewriter:
    def __init__(self, api_type: str = "tongyi"):
        """
        初始化 ContentRewriter
        :param api_type: API类型，可选值："zhipu" 或 "tongyi"
        """
        self.extractor = DocumentExtractor()
        # 根据api_type选择API实现
        if api_type.lower() == "tongyi":
            from tongyi_api import TongYiAPI
            self.ai = TongYiAPI()
        else:  # 默认使用 zhipu
            from zhipu_api import ZhipuAPI
            self.ai = ZhipuAPI()
        
        # 动态计算最佳线程数: CPU核心数 * 2（但不超过系统内存和API限制）
        cpu_count = multiprocessing.cpu_count()
        mem = psutil.virtual_memory()
        # 根据可用内存估算最大线程数：每100MB内存支持1个线程（经验值）
        mem_threads = max(1, int(mem.available / (100 * 1024 * 1024)))
        # 取较小值作为实际线程数，并限制最大值为50以避免API限流
        self.max_workers = min(cpu_count * 2, mem_threads, 50)
        logger.info(f"设置线程池大小为: {self.max_workers} (CPU核心数: {cpu_count}, 内存支持线程数: {mem_threads})")
        
        # 设置缓存目录
        self.base_output_dir = os.path.abspath("output")
        self.cache_dir = os.path.join(self.base_output_dir, "cache")
        self.images_dir = os.path.join(self.base_output_dir, "images")
        self.docx_dir = os.path.join(self.base_output_dir, "docx_files")
        self.markdown_dir = os.path.join(self.base_output_dir, "markdown_files")
        
        # 创建必要的目录
        for dir_path in [self.cache_dir, self.images_dir, self.docx_dir, self.markdown_dir]:
            os.makedirs(dir_path, exist_ok=True)
        
        # 加载缓存（如果存在）
        self.cache = self._load_cache()
        
        # 批处理大小：用于分批处理大型文档
        self.batch_size = 10
        
        # 进度阶段权重分配 (总和为100)
        self.progress_weights = {
            'init': 2,           # 初始化阶段
            'extraction': 15,    # 文档内容提取阶段
            'rewrite': 60,       # 内容重写阶段 (最耗时)
            'process_non_text': 5, # 处理非文本内容 (图片、公式等)
            'generate_docx': 8,   # 生成Word文档
            'generate_markdown': 7, # 生成Markdown文档
            'finalize': 3        # 最终清理和完成工作
        }

    def _load_cache(self) -> Dict[str, str]:
        """加载文本重写缓存"""
        cache_file = os.path.join(self.cache_dir, "rewrite_cache.json")
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"加载缓存失败: {str(e)}，将使用空缓存")
        return {}
    
    def _save_cache(self) -> None:
        """保存文本重写缓存"""
        cache_file = os.path.join(self.cache_dir, "rewrite_cache.json")
        try:
            # 限制缓存大小（最多保存1000条）
            if len(self.cache) > 1000:
                # 仅保留最近的1000个条目
                self.cache = dict(list(self.cache.items())[-1000:])
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"保存缓存失败: {str(e)}")

    def _get_text_hash(self, text: str) -> str:
        """生成文本的哈希值作为缓存键"""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def _update_progress(self, 
                         stage: str, 
                         stage_progress: float, 
                         progress_callback: Optional[Callable[[int, int, Optional[str]], None]], 
                         message: Optional[str] = None) -> None:
        """
        更新总体进度
        
        Args:
            stage: 当前阶段名称 (与self.progress_weights中的键对应)
            stage_progress: 当前阶段的进度 (0.0-1.0)
            progress_callback: 进度回调函数
            message: 可选的进度消息
        """
        if not progress_callback:
            return
        
        # 计算当前阶段前的累积权重
        accumulated_weight = 0
        for s, w in self.progress_weights.items():
            if s == stage:
                break
            accumulated_weight += w
        
        # 计算当前阶段的进度贡献
        stage_contribution = self.progress_weights[stage] * stage_progress
        
        # 计算总体进度
        total_progress = int(accumulated_weight + stage_contribution)
        total_progress = min(100, max(0, total_progress))  # 确保在0-100范围内
        
        # 调用回调函数更新进度
        message_text = message if message else f"阶段: {stage} ({int(stage_progress * 100)}%)"
        progress_callback(total_progress, 100, message_text)

    def rewrite_content(self, input_file: str, progress_callback: Optional[Callable[[int, int, Optional[str]], None]] = None) -> None:
        """
        重写文档内容
        
        Args:
            input_file: 输入文件路径
            progress_callback: 进度回调函数，接受当前进度、总进度和可选的消息参数
        """
        start_time = time.time()
        logger.info(f"开始处理文件: {input_file}")
        
        # 初始化阶段 (2%)
        self._update_progress('init', 0.5, progress_callback, "正在初始化...")
        
        # 生成唯一的文档名
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        base_filename = os.path.splitext(os.path.basename(input_file))[0]
        output_filename = f"{base_filename}_{timestamp}"

        self._update_progress('init', 1.0, progress_callback, "初始化完成")
        
        # 内容提取阶段 (15%)
        self._update_progress('extraction', 0.1, progress_callback, f"正在提取文档内容: {base_filename}...")

        # 提取内容
        extract_start = time.time()
        content = self.extractor.extract_content(input_file)
        extract_time = time.time() - extract_start
        
        total_blocks = len(content['content_blocks'])
        logger.info(f"共发现 {total_blocks} 个内容块，耗时: {extract_time:.2f}秒")
        
        # 提取完成的进度更新
        self._update_progress('extraction', 1.0, progress_callback, 
                             f"提取完成，共发现 {total_blocks} 个内容块，耗时: {extract_time:.2f}秒")

        # 计算需要重写的文本块
        text_blocks = [block for block in content['content_blocks'] if block['type'] == 'text']
        text_blocks_count = len(text_blocks)
        processed_text_blocks = 0
        
        # 内容重写阶段 (60%)
        self._update_progress('rewrite', 0.0, progress_callback, 
                             f"开始重写内容，共有 {text_blocks_count} 个文本块...")
        
        # 分批处理文本块以提高效率
        for i in range(0, len(text_blocks), self.batch_size):
            batch = text_blocks[i:i+self.batch_size]
            batch_size = len(batch)
            
            # 更新批次开始信息
            batch_num = i // self.batch_size + 1
            total_batches = (len(text_blocks) + self.batch_size - 1) // self.batch_size
            self._update_progress('rewrite', processed_text_blocks / max(1, text_blocks_count), 
                                 progress_callback, f"处理批次 {batch_num}/{total_batches}, 批次大小: {batch_size}")
            
            # 使用线程池并行处理当前批次的文本块
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # 创建任务列表
                future_to_block = {}
                for block in batch:
                    future = executor.submit(self._process_text_block, block)
                    future_to_block[future] = block

                # 处理完成的任务
                for future in as_completed(future_to_block):
                    block = future_to_block[future]
                    try:
                        block['content'] = future.result()
                        processed_text_blocks += 1
                        
                        # 更新重写阶段的进度
                        stage_progress = processed_text_blocks / max(1, text_blocks_count)
                        self._update_progress('rewrite', stage_progress, progress_callback, 
                                            f"重写进度: {processed_text_blocks}/{text_blocks_count}")
                    except Exception as e:
                        logger.error(f"处理文本块时出错: {str(e)}")
        
        # 确保重写阶段的最终进度为100%
        self._update_progress('rewrite', 1.0, progress_callback, 
                             f"文本重写完成，共处理了 {processed_text_blocks} 个文本块")
        
        # 处理非文本内容 (5%)
        non_text_blocks = [block for block in content['content_blocks'] if block['type'] != 'text']
        non_text_count = len(non_text_blocks)
        
        self._update_progress('process_non_text', 0.0, progress_callback, 
                             f"开始处理非文本内容，共有 {non_text_count} 个项目...")
        
        # 处理非文本块
        for i, block in enumerate(non_text_blocks):
            self._process_non_text_block(block, output_filename)
            # 更新非文本处理进度
            self._update_progress('process_non_text', (i + 1) / max(1, non_text_count), 
                                 progress_callback, f"处理非文本内容: {i + 1}/{non_text_count}")

        # 确保非文本处理阶段的最终进度为100%
        self._update_progress('process_non_text', 1.0, progress_callback, "非文本内容处理完成")

        logger.info("\n所有内容块处理完成，开始生成文档...")
        
        # 生成DOCX文档 (8%)
        self._update_progress('generate_docx', 0.0, progress_callback, "开始生成Word文档...")
        
        try:
            # 生成Word文档
            docx_start = time.time()
            docx_path = self._generate_docx(content['content_blocks'], output_filename)
            docx_time = time.time() - docx_start
            
            self._update_progress('generate_docx', 1.0, progress_callback, 
                                f"Word文档生成完成，耗时: {docx_time:.2f}秒")
                
            # 生成Markdown文档 (7%)
            self._update_progress('generate_markdown', 0.0, progress_callback, "开始生成Markdown文档...")
            
            markdown_start = time.time()
            self._generate_markdown(content['content_blocks'], output_filename)
            markdown_time = time.time() - markdown_start
            
            self._update_progress('generate_markdown', 1.0, progress_callback, 
                                 f"Markdown文档生成完成，耗时: {markdown_time:.2f}秒")
            
            # 最终阶段 (3%)
            self._update_progress('finalize', 0.0, progress_callback, "开始最终处理...")
            
            # 保存缓存
            self._save_cache()
            
            # 计算总处理时间
            elapsed_time = time.time() - start_time
            logger.info(f"\n文档处理完成！总耗时: {elapsed_time:.2f}秒")
            logger.info(f"生成的文件位置：{docx_path}")
            
            # 最终完成
            self._update_progress('finalize', 1.0, progress_callback, 
                                 f"处理完成！总耗时: {elapsed_time:.2f}秒")
        finally:
            # 确保所有文档生成完成后才清理临时图片
            self._cleanup_images()

    def _process_text_block(self, block: Dict) -> str:
        """处理文本块，使用缓存提高效率"""
        # 检查是否包含公式标记，如果包含则跳过重写
        if self._contains_formula(block['content']):
            logger.info(f"检测到公式，跳过重写: {block['content'][:30]}...")
            return block['content']
        
        # 计算文本哈希作为缓存键
        text_hash = self._get_text_hash(block['content'])
        
        # 检查缓存
        if text_hash in self.cache:
            logger.info(f"使用缓存的重写结果: {block['content'][:30]}...")
            return self.cache[text_hash]
        
        # 缓存未命中，调用API重写
        try:
            result = self.ai.rewrite_text(block['content'])
            # 保存到缓存
            self.cache[text_hash] = result
            return result
        except Exception as e:
            logger.error(f"调用AI API失败: {str(e)}")
            # 出错时返回原文本，确保不影响整体流程
            return block['content']

    def _contains_formula(self, text: str) -> bool:
        """检查文本中是否包含公式标记"""
        # LaTeX公式标记
        latex_markers = ['\\begin{equation}', '\\end{equation}', '\\begin{align}', '\\end{align}', 
                         '$', '\\frac', '\\sum', '\\int', '\\alpha', '\\beta', '\\gamma']
        
        # Office数学公式可能包含的XML标记
        office_markers = ['<m:oMath', '<m:r>', '<m:t>', '<m:f>', '<m:num>', '<m:den>']
        
        # 检查是否包含任何公式标记
        for marker in latex_markers + office_markers:
            if marker in text:
                return True
        
        return False

    def _copy_file_with_retry(self, src: str, dst: str, max_retries: int = 3, delay: float = 1.0) -> bool:
        """带重试机制的文件复制
        
        Args:
            src: 源文件路径
            dst: 目标文件路径
            max_retries: 最大重试次数
            delay: 重试间隔（秒）
            
        Returns:
            bool: 是否复制成功
        """
        for attempt in range(max_retries):
            try:
                # 确保目标目录存在
                os.makedirs(os.path.dirname(dst), exist_ok=True)
                
                # 如果目标文件已存在，先尝试删除
                if os.path.exists(dst):
                    try:
                        os.remove(dst)
                    except:
                        pass
                
                # 复制文件
                shutil.copy2(src, dst)
                return True
            except Exception as e:
                logger.warning(f"文件复制失败 (尝试 {attempt + 1}/{max_retries}): {str(e)}")
                if attempt < max_retries - 1:
                    time.sleep(delay)
                continue
        return False

    def _process_non_text_block(self, block: Dict, output_filename: str) -> None:
        """处理非文本块（图片、标题等）"""
        # 处理公式块（如果有特殊标记）
        if block.get('is_formula', False):
            logger.info(f"处理公式块: {block.get('content', '')[:30]}...")
            # 保留原始公式，不做修改
            pass
        
        # 图片已经在DocumentExtractor中保存到output/images目录，这里不需要再复制
        pass

    def _generate_docx(self, blocks: List[Dict], output_filename: str) -> str:
        """生成DOCX文档，优化批量处理"""
        doc = Document()
        
        for block in blocks:
            if block['type'] == 'text':
                para = doc.add_paragraph(block['content'])
                # 检查是否有格式信息
                if 'format_info' in block:
                    self._apply_format_to_paragraph(para, block['format_info'])
            elif block['type'] == 'heading':
                heading = doc.add_heading(block['content'], block.get('level', 1))
                # 检查是否有格式信息
                if 'format_info' in block:
                    self._apply_format_to_paragraph(heading, block['format_info'])
            elif block['type'] == 'image':
                if 'image_path' in block:
                    try:
                        logger.info(f"正在插入图片: {os.path.basename(block['image_path'])}")
                        doc.add_paragraph()  # 空行
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(block['image_path'], width=Inches(6))
                        doc.add_paragraph()  # 空行
                        logger.info(f"✓ 图片插入成功: {block['image_path']}")
                    except Exception as e:
                        logger.error(f"✗ 图片插入失败: {str(e)}")
                else:
                    logger.warning("✗ 警告: 无法插入图片，缺少图片路径")
            elif block['type'] == 'formula':
                # 对于公式类型，添加特殊处理
                try:
                    para = doc.add_paragraph(block['content'])
                    if 'format_info' in block:
                        self._apply_format_to_paragraph(para, block['format_info'])
                    logger.info(f"✓ 公式插入成功")
                except Exception as e:
                    logger.error(f"✗ 公式插入失败: {str(e)}")
        
        output_path = os.path.join(self.docx_dir, f"{output_filename}.docx")
        doc.save(output_path)
        logger.info(f"DOCX文档已生成: {output_path}")
        return output_path
    
    def _apply_format_to_paragraph(self, paragraph, format_info: Dict) -> None:
        """应用格式信息到段落"""
        try:
            if 'alignment' in format_info and format_info['alignment'] != 'None':
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                alignment_map = {
                    '0': WD_ALIGN_PARAGRAPH.LEFT,
                    '1': WD_ALIGN_PARAGRAPH.CENTER,
                    '2': WD_ALIGN_PARAGRAPH.RIGHT,
                    '3': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if format_info['alignment'] in alignment_map:
                    paragraph.alignment = alignment_map[format_info['alignment']]
            
            # 设置首行缩进
            if 'first_line_indent' in format_info and format_info['first_line_indent'] is not None:
                paragraph.paragraph_format.first_line_indent = format_info['first_line_indent']
            
            # 设置行距
            if 'line_spacing' in format_info and format_info['line_spacing'] is not None:
                paragraph.paragraph_format.line_spacing = format_info['line_spacing']
        except Exception as e:
            logger.warning(f"应用格式失败: {str(e)}")

    def _generate_markdown(self, blocks: List[Dict], output_filename: str) -> None:
        """生成Markdown文档，优化批量处理"""
        markdown_content = []
        
        # 复制图片到markdown专用的图片目录
        markdown_images_dir = os.path.join(self.markdown_dir, "images")
        os.makedirs(markdown_images_dir, exist_ok=True)
        
        # 一次性收集所有图片复制任务
        copy_tasks = []
        
        for block in blocks:
            if block['type'] == 'text':
                markdown_content.append(block['content'] + "\n\n")
            elif block['type'] == 'heading':
                level = block.get('level', 1)
                markdown_content.append('#' * level + ' ' + block['content'] + "\n\n")
            elif block['type'] == 'formula':
                # 对于公式，使用Markdown的数学语法
                # 使用两个$$包围的是独立公式块，使用一个$包围的是行内公式
                markdown_content.append(f"\n$${block['content']}$$\n\n")
            elif block['type'] == 'image':
                if 'image_path' in block:
                    image_filename = os.path.basename(block['image_path'])
                    markdown_image_path = os.path.join(markdown_images_dir, image_filename)
                    
                    # 添加到复制任务列表
                    copy_tasks.append((block['image_path'], markdown_image_path, image_filename))
        
        # 并行执行图片复制任务
        with ThreadPoolExecutor(max_workers=min(10, self.max_workers)) as executor:
            future_to_task = {executor.submit(self._copy_file_with_retry, src, dst): (src, dst, filename) 
                             for src, dst, filename in copy_tasks}
            
            for future in as_completed(future_to_task):
                src, dst, filename = future_to_task[future]
                try:
                    success = future.result()
                    if success:
                        # 使用相对路径引用图片
                        index = next((i for i, task in enumerate(copy_tasks) if task[2] == filename), -1)
                        if index >= 0:
                            markdown_content.insert(index, f"\n![{filename}](./images/{filename})\n\n")
                        logger.info(f"✓ Markdown图片复制成功: {filename}")
                    else:
                        logger.error(f"✗ Markdown图片复制失败: {filename}")
                except Exception as e:
                    logger.error(f"✗ Markdown图片处理出错: {str(e)}")
        
        # 保存markdown文件
        output_path = os.path.join(self.markdown_dir, f"{output_filename}.md")
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(''.join(markdown_content))
            logger.info(f"✓ Markdown文档已生成: {output_path}")
        except Exception as e:
            logger.error(f"✗ Markdown文档生成失败: {str(e)}")

    def _cleanup_images(self) -> None:
        """清理临时图片文件"""
        try:
            # 只清理主图片目录，保留markdown专用的图片
            if os.path.exists(self.images_dir):
                # 使用重试机制删除目录
                for attempt in range(3):
                    try:
                        shutil.rmtree(self.images_dir)
                        logger.info("✓ 已清理临时图片文件")
                        break
                    except Exception as e:
                        if attempt == 2:  # 最后一次尝试
                            logger.warning(f"✗ 清理临时图片文件失败: {str(e)}")
                        else:
                            time.sleep(1)
                            continue
        except Exception as e:
            logger.warning(f"✗ 清理临时图片文件时出错: {str(e)}")